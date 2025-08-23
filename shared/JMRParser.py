import os, sys, logging, time, re, json
import requests
import zipfile
import openpyxl, time
from openpyxl.styles import Font
from typing import List, Optional
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from pykakasi import kakasi


kks = kakasi()

# Mapping of counters with their special readings
COUNTER_MAPPINGS = {
    "人": {
        1: "{一人|ひとり}",
        2: "{二人|ふたり}",
        3: "{三人|さんにん}",
        4: "{四人|よにん}",
        5: "{五人|ごにん}",
        6: "{六人|ろくにん}",
        7: "{七人|ななにん}",
        8: "{八人|はちにん}",
        9: "{九人|きゅうにん}",
        10: "{十人|じゅうにん}",
    }
}
# prefer きみ over くん for standalone 君
PREFERRED_READING = {
    "君": "きみ",   
}

def fs_safe(name: str) -> str:
    # replace forbidden chars with "_", and avoid trailing dot/space (Windows)
    name = re.sub(r'[<>:"/\\|?*`#^\[\]]+', "_", name)
    return name.rstrip(" .") or "_"

_COUNTER_PATTERN = re.compile(
    r"(\d+)(" + "|".join(map(re.escape, COUNTER_MAPPINGS.keys())) + r")"
)

skipped_line_message = [None]
log_path = os.path.join(os.path.dirname(sys.argv[0]), "furigana_parser.log")
logging.basicConfig(
    filename=log_path,
    level=logging.DEBUG,
    format="%(asctime)s %(levelname)s: %(message)s",
    encoding="utf-8"
)
def get_metadata_from_zip(model_path):
    with zipfile.ZipFile(model_path, 'r') as zip_file:
        with zip_file.open('package.argosmodel') as model_file:
            return json.loads(model_file.read().decode('utf-8'))
        
def heavy_initialization():
    global translator
    import argostranslate.package
    import argostranslate.translate
    if hasattr(sys, '_MEIPASS'):
        base_path = sys._MEIPASS
    else:
        sys._MEIPASS = os.path.abspath("test_bundle_root")
        base_path = sys._MEIPASS

    model_path = os.path.join(base_path, "argos-translate", "packages", "ja_en.zip")

    if os.path.exists(model_path):
        metadata = get_metadata_from_zip(model_path)
        from_code = metadata['from_code']
        to_code = metadata['to_code']
        installed = any(
        pkg.from_code == from_code and pkg.to_code == to_code
        for pkg in argostranslate.translate.get_installed_packages()
        )
        if not installed:
            argostranslate.package.install_from_path(model_path)
    else:
        print(f"Model not found: {model_path}")

    installed_languages = argostranslate.translate.get_installed_languages()
    ja = next((lang for lang in installed_languages if lang.code == "ja"), None)
    en = next((lang for lang in installed_languages if lang.code == "en"), None)
    translator = ja.get_translation(en) if ja and en else None

def is_kanji(char: str) -> bool:
    return '一' <= char <= '龯'

def is_katakana(text: str) -> bool:
    """Return True if the entire string consists of Katakana characters."""
    return bool(text) and all('\u30A0' <= ch <= '\u30FF' for ch in text)

def convert_line_to_ruby_pairs(line: str):
    def _replace_counters(text: str) -> str:
        def _repl(match: re.Match) -> str:
            num = int(match.group(1))
            counter = match.group(2)
            reading = COUNTER_MAPPINGS.get(counter, {}).get(num)
            return reading if reading else match.group(0)
        return _COUNTER_PATTERN.sub(_repl, text)

    line = _replace_counters(line)
    result = []
    for item in kks.convert(line):
        orig = item['orig']
        hira = item['hira']

        # Katakana → keep as-is (no ruby) and don't process it again
        if is_katakana(orig):
            result.append((orig, None))
            continue

        # If the token has no kanji, leave it as-is
        if not any(is_kanji(c) for c in orig):
            result.append((orig, None))
            continue

        # If the token is entirely kanji, keep the existing mapping
        if all(is_kanji(c) for c in orig):
            result.append((orig, hira))
            continue

        # Token has a mix of kanji and kana/punctuation
        idx = 0
        chars = list(orig)
        i = 0
        while i < len(chars):
            ch = chars[i]
            if is_kanji(ch):
                # Determine contiguous kanji run
                j = i
                while j < len(chars) and is_kanji(chars[j]):
                    j += 1

                # Find reading boundary using the next non-kanji character
                if j < len(chars):
                    next_hira = kks.convert(chars[j])[0]["hira"]
                    next_idx = hira.find(next_hira, idx)
                    if next_idx == -1:
                        next_idx = len(hira)
                else:
                    next_idx = len(hira)

                kanji_reading = hira[idx:next_idx]
                run = chars[i:j]

                if len(kanji_reading) == len(run):
                    for k, kc in enumerate(run):
                        result.append((kc, kanji_reading[k]))
                else:
                    remaining = kanji_reading
                    for k, kc in enumerate(run):
                        if k == len(run) - 1:
                            reading = remaining
                        else:
                            guess = kks.convert(kc)[0]["hira"]
                            if remaining.startswith(guess):
                                reading = guess
                                if remaining:
                                    remaining = remaining[len(reading):]
                            else:
                                if remaining:
                                    reading = remaining[0]
                                    remaining = remaining[1:]
                                else:
                                    reading = guess or ""
                        result.append((kc, reading))

                idx = next_idx
                i = j
            else:
                ch_hira = kks.convert(ch)[0]["hira"]
                result.append((ch, None))
                idx += len(ch_hira)
                i += 1

    # -------- Preferred reading overrides ---------------------------------

    def _is_name_suffix_context(prev_base: str) -> bool:
        """Return True when previous token looks like a name part (so '君' is likely a suffix)."""
        if not prev_base:
            return False
        # Japanese letters
        if any('\u4E00' <= c <= '\u9FFF' or '\u3040' <= c <= '\u30FF' for c in prev_base):
            return True
        # ASCII letters/digits or middle dot / long vowel mark often used in names
        if re.search(r'[A-Za-z0-9]', prev_base) or '・' in prev_base or 'ー' in prev_base:
            return True
        return False

    for n, (base, reading) in enumerate(result):
        pref = PREFERRED_READING.get(base)
        if not pref or not reading:
            continue

        # Normalize entry: accept "reading" or ("reading","mode")
        if isinstance(pref, tuple):
            pref_reading, mode = pref[0], (pref[1] if len(pref) > 1 else "always")
        else:
            pref_reading, mode = pref, "always"

        if mode == "standalone":
            prev_base = result[n-1][0] if n > 0 else ""
            if _is_name_suffix_context(prev_base):
                continue  # keep original reading (likely ～君 suffix)

        # Apply override
        result[n] = (base, pref_reading)

    return result

def add_ruby_eq_field(paragraph, base_text, ruby_text, base_font_size_pt=16):
    # Create the run that will hold the field
    run_instr = paragraph.add_run()
    
    # Set the font size of the run (this controls base text size)
    run_instr.font.size = Pt(base_font_size_pt)
    
    fldChar_begin = OxmlElement('w:fldChar')
    fldChar_begin.set(qn('w:fldCharType'), 'begin')
    run_instr._r.append(fldChar_begin)

    instr_text = OxmlElement('w:instrText')
    instr_text.set(qn('xml:space'), 'preserve')

    # Furigana size in half-points (hps18 = 9pt), adjust if needed
    furigana_hps = 18  

    instr_text.text = (
        f'EQ \\* jc2 \\* "Font:Noto Sans JP Light" \\* hps{furigana_hps} '
        f'\\o\\ad(\\s\\up 17({ruby_text}),{base_text})'
    )
    run_instr._r.append(instr_text)

    fldChar_end = OxmlElement('w:fldChar')
    fldChar_end.set(qn('w:fldCharType'), 'end')
    run_instr._r.append(fldChar_end)

def generate_obsidian_lyric_file(
    lyrics_lines: List[str],
    song_title: str,           # display title (unsanitized; e.g., "AM6:30")
    artist: str,
    album: str,
    track_number: int,         # album number for this song (e.g., 2)
    total_tracks: int,         # highest album number
    track_titles: List[str],   # raw titles (may include numbers)
    output_root: str = "Lyrics"
):
    # 0) Skip if there are no real lyrics
    if not lyrics_lines or not any((ln or "").strip() for ln in lyrics_lines):
        return None

    # Helpers (local so this function is self-contained)
    TRACK_NO_RE = re.compile(r'^\s*(\d{1,3})\s*[\.．]?\s*')
    def parse_track_no(title: str) -> int | None:
        m = TRACK_NO_RE.match(title or "")
        return int(m.group(1)) if m else None
    def strip_track_prefix(title: str) -> str:
        return TRACK_NO_RE.sub("", title or "").strip()
    def sanitize_filename(text: str) -> str:
        # safe for Windows/macOS/Linux
        text = re.sub(r'[<>:"/\\|?*`#^\[\]]+', "_", text or "")
        return text.rstrip(" .") or "_"

    def line_to_furigana(line: str) -> str:
        pairs = convert_line_to_ruby_pairs(line)
        return ''.join(
            f'{{{base}|{reading}}}' if reading and base != reading else base
            for base, reading in pairs
        )

    # 1) Paths
    folder_path = os.path.join(
        output_root, "Lyrics",
        sanitize_filename(artist),
        sanitize_filename(album),
    )
    os.makedirs(folder_path, exist_ok=True)

    filename = f"{track_number:02d}. {sanitize_filename(song_title)}.md"
    file_path = os.path.join(folder_path, filename)

    # 2) Build map: album number -> plain title (for accurate prev/next)
    num_to_title = {}
    for t in track_titles:
        n = parse_track_no(t)
        if n:  # only map those we can number
            num_to_title.setdefault(n, strip_track_prefix(t))

    prev_num = track_number - 1 if track_number > 1 else None
    next_num = track_number + 1 if track_number < total_tracks else None

    previous_filename = (
        f"{prev_num:02d}. {sanitize_filename(num_to_title[prev_num])}"
        if (prev_num and prev_num in num_to_title) else None
    )
    next_filename = (
        f"{next_num:02d}. {sanitize_filename(num_to_title[next_num])}"
        if (next_num and next_num in num_to_title) else None
    )

    # 3) Write file
    with open(file_path, "w", encoding="utf-8") as f:
        f.write("---\n")
        f.write(f"title: {song_title}\n")     # display title (unsanitized)
        f.write(f"artist: {artist}\n")
        f.write(f"album: {album}\n")
        f.write(f"track: {track_number}\n")
        f.write("tags: [lyrics, japanese, furigana]\n")
        f.write("language: ja\n")
        f.write("---\n\n")

        if previous_filename:
            f.write(f"← [[{previous_filename}]]\n")
        f.write("[[link]]\n\n")

        for line in lyrics_lines:
            s = (line or "").strip()
            f.write((line_to_furigana(s) if s else "") + "\n")

        f.write("\n")
        if next_filename:
            f.write(f"[[{next_filename}]] →\n")
        f.write("\n[[link]]\n")

    return file_path


def create_docx_with_eq_fields(input_path, output_path):
    document = Document()
    # Safely get and configure 'Normal' style
    style = document.styles['Normal']
    style.font.name = 'Noto Sans JP Light'

    # Safely set east Asian font
    rPr = style._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), 'Noto Sans JP Light')

    # Set paragraph spacing (single line, no extra space after)
    paragraph_format = style.paragraph_format
    paragraph_format.line_spacing = 1.0
    paragraph_format.space_after = Pt(0)

    with open(input_path, encoding='utf-8') as f:
        lines = f.readlines() 

    for line in lines:
        p = document.add_paragraph()
        
        pairs = convert_line_to_ruby_pairs(line.rstrip("\n"))

        for base, reading in pairs:
            if reading:
                add_ruby_eq_field(p, base, reading)
            else:
                run = p.add_run(base)
                run.font.size = Pt(16)
                run.font.name = 'Noto Sans JP Light'
                r = run._element.rPr
                rFonts = OxmlElement('w:rFonts')
                rFonts.set(qn('w:eastAsia'), 'Noto Sans JP Light')
                r.append(rFonts)

        # insert a line break **within** the same paragraph
        # p.add_run().add_break()

    document.save(output_path)

skipped_line_number = [None]  # mutable container to store the first skipped line number

def ui_warning_callback(msg):
    if skipped_line_message[0] is None:
        skipped_line_message[0] = msg

def process_lines_with_options(
    input_path: str,
    output_path: str,
    manual_xlsx: Optional[str] = None,
    use_offline: bool = True,
    use_online: bool = False,
    export_spreadsheet: bool = False,
    progress_callback=None,
    ui_warning_callback=None
):
    manual_translations = load_manual_translation(manual_xlsx) if manual_xlsx else {}
    with open(input_path, "r", encoding="utf-8") as f:
        lines = f.readlines()

    spreadsheet_data = []
    output_data = []

    total = len(lines)

    for i, line in enumerate(lines):
        clean_line = line.strip() 

        if clean_line == "":
            # This is a blank line
            styled_jp = ""   # For spreadsheet and Word, keep blank
            manual = ""
            local = ""
            online = ""

            # For JSON output, use paragraph break symbol instead of empty string
            json_jp_text = "\n\n"

        else:
            pairs = convert_line_to_ruby_pairs(clean_line)
            styled_jp = "".join(
                [f"<ruby={reading}>{base}</ruby>" if reading else base for base, reading in pairs]
            )
            manual = manual_translations.get(clean_line, "")
            if manual_xlsx is not None and manual == "":
                warning_msg = f"Manual translation skipped for line: '{clean_line[:30]}...'"
                if ui_warning_callback:
                    ui_warning_callback(f"Partial manual translation: Japanese doesn't match from Row {i + 2}")
            local = translator.translate(clean_line) if translator and use_offline else ""
            online = translate_online(clean_line) if use_online else ""
            json_jp_text = styled_jp  # normal line styled for JSON

        spreadsheet_data.append([clean_line, manual, local, online])
        print(f"Line {i}: '{clean_line}' (empty? {clean_line == ''})")
        entry = {"jp_text": json_jp_text}

        if manual_xlsx is not None:
            entry["manual"] = manual or ""

        if use_offline:
            entry["local"] = local or ""

        if use_online:
            entry["online"] = online or ""

        output_data.append(entry)

        if progress_callback:
            progress_callback(i, total)

    with open(output_path, "w", encoding="utf-8") as f:
        import json
        json.dump(output_data, f, ensure_ascii=False, indent=2)

    if export_spreadsheet:
        xlsx_path = os.path.splitext(output_path)[0] + ".xlsx"
        save_spreadsheet(
            xlsx_path,
            spreadsheet_data,
            include_manual=bool(manual_xlsx),
            include_local=use_offline,
            include_online=use_online
        )

def load_manual_translation(path: str) -> dict:
    wb = openpyxl.load_workbook(path)
    sheet = wb.active
    data = {}
    for row in sheet.iter_rows(min_row=2):
        jp = row[0].value
        manual = row[1].value
        if jp and manual:
            data[jp.strip()] = manual.strip()
    return data

def save_spreadsheet(output_path: str, rows: List[List[str]], include_manual=True, include_local=True, include_online=True):
    wb = openpyxl.Workbook()
    ws = wb.active
    
    headers = ["Japanese"]
    if include_manual:
        headers.append("Manual Translation")
    if include_local:
        headers.append("Local Translation")
    if include_online:
        headers.append("Online Translation")

    ws.append(headers)
    for cell in ws[1]:
        cell.font = Font(bold=True)
    
    for row in rows:
        filtered_row = [row[0]]
        if include_manual:
            filtered_row.append(row[1])
        if include_local:
            filtered_row.append(row[2])
        if include_online:
            filtered_row.append(row[3])
        ws.append(filtered_row)

    wb.save(output_path)

def translate_online(text: str) -> str:
    try:
        url = "https://libretranslate.de/translate"
        payload = {
            "q": text,
            "source": "ja",
            "target": "en",
            "format": "text"
        }
        headers = {
            "Content-Type": "application/json"
        }
        response = requests.post(url, json=payload, headers=headers, timeout=10)
        time.sleep(3)  # wait 3 seconds to avoid rate limiting
        if response.status_code == 200:
            result = response.json()
            return result.get("translatedText", "")
        else:
            return ""
    except Exception as e:
        return ""